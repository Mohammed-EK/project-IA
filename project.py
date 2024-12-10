import streamlit as st
import praw
import re
from docx import Document
import io
from time import sleep
from stqdm import stqdm  # For a progress bar in Streamlit
import matplotlib.pyplot as plt
from matplotlib.ticker import MaxNLocator
import docx.shared

# Setting up the Reddit API using PRAW
reddit = praw.Reddit(
    client_id="41Huqnmhw4KOlpZth6nAWQ",
    client_secret="GPt_YaaJ4Iyv9ZVmyXFkerO0nCdfkQ",
    user_agent="AggressiveCommentDetector/1.0 by u/Separate-Donkey-7902"
)

# Aggressive keywords in English and French
AGGRESSIVE_KEYWORDS = [
    # English
    "hate", "stupid", "idiot", "moron", "dumb", "shut up", "kill", "worthless",
    "trash", "fool", "ignorant", "pathetic", "useless", "annoying", "disgusting",
    "loser", "jerk", "failure", "coward", "weak", "liar", "bully", "scum", 
    "pest", "corrupt", "thief", "rage", "revenge", "burn", "damn", "hell", 
    "fuck", "fucking", "shit", "bullshit", "bitch", "cunt", "dickhead", "motherfucker",
    "goddamn", "freaking", "asshole", "piss off", "crap", "bastard", "prick", 
    "whore", "slut", "fuckface", "screw you", "dumbass", "retard", "scumbag", 
    "nitwit", "dipshit", "tool", "half-wit", "degenerate", "vermin", "parasite", 
    "low-life", "clown", "wannabe", "delusional", "piece of shit", "loser", "scrub", 
    "noob", "rage quit", "crybaby", "salty", "get rekt", "gremlin", "cheater", 
    "hacker", "toxic", "worthless piece of shit", "filthy animal", "die", "rot", 
    "drop dead", "go to hell", "burn in hell", "eat shit", "choke", "fuck off",
    
    # French
    "haine", "stupide", "idiot", "imbécile", "nul", "dégage", "inutile", 
    "déchet", "abruti", "ignorant", "pathétique", "dégoûtant", "perdant", 
    "salaud", "lâche", "menteur", "faible", "meurtrier", "destructeur", 
    "pourri", "merde", "enfer", "enfoiré", "connard", "casse-toi", "putain", 
    "bordel", "salope", "enculé", "nique ta mère", "ferme ta gueule", "chiotte", 
    "pédé", "con", "connasse", "connard", "foutre", "va te faire foutre", 
    "va crever", "pourrir", "chier dessus", "abrutissement", "ordure", "raté", 
    "baltringue", "clodo", "mongol", "rageux", "croûton", "bouffon", 
    "branleur", "crétin", "plouc", "tocard", "kéké", "espèce de merde", 
    "sale chien", "batard", "fils de pute", "foutaises",
    
    # Aggressive Phrases
    "you suck", "nobody likes you", "die already", "go die", "kill yourself", 
    "shut the fuck up", "get lost", "waste of space", "go screw yourself", 
    "drown", "burn", "choke on it", "eat dirt", "filthy rat", "you’re pathetic",
    
    # Internet-Specific
    "git gud", "ez clap", "tryhard", "ragequit", "go touch grass", "malding", 
    "seethe", "cope", "L + ratio", "owned", "rekt", "trash player", "camping noob", 
    "spawn camper", "lagging cheater", "greedy loot goblin", "report him", 
    "bot", "spammer", "troll", "griefer", "toxic gamer",

    # Profanity Variations
    "fuck", "f***","fk", "shit", "sh**", "bitch", "b*tch", "cunt", "c*nt", 
    "asshole", "a**hole", "dickhead", "d**khead", "motherfucker", "motherf***er", 
    "goddamn", "god d*mn", "fucking", "f*cking", "ass", "a*s", "crap", "s***", 
    "bullshit", "bulls***", "whore", "w***e", "slut", "sl**", "prick", "pr**k"
]

# Function to analyze a user's comments for aggressive behavior
def fetch_aggressive_comments_by_user(username):
    try:
        user = reddit.redditor(username)
        aggressive_comments = []
        total_comments = 0

        posts = list(user.submissions.new(limit=10))
        total_posts = len(posts)

        with stqdm(posts, desc="Fetching posts...", total=total_posts) as progress_bar:
            for post in progress_bar:
                post.comments.replace_more(limit=0)
                for comment in post.comments.list():
                    total_comments += 1
                    sentiment_score = 0

                    # Calculate sentiment score
                    for keyword in AGGRESSIVE_KEYWORDS:
                        if keyword.lower() in comment.body.lower():
                            sentiment_score -= 0.5 if keyword.islower() else -0.75
                    sentiment_score *= 1.5

                    if sentiment_score < 0:
                        aggressive_comments.append({
                            'post_title': post.title,
                            'comment_body': comment.body,
                            'sentiment_score': sentiment_score
                        })

        return aggressive_comments, total_comments
    except Exception as e:
        st.error(f"An error occurred: {e}")
        return [], 0

# Function to create a cumulative chart image
def create_cucumber_chart(aggressive_comments):
    scores = [comment['sentiment_score'] for comment in aggressive_comments]
    scores.sort()
    cumulative_scores = [sum(scores[:i+1]) for i in range(len(scores))]

    fig, ax = plt.subplots()
    ax.plot(range(1, len(cumulative_scores) + 1), cumulative_scores, marker='o', color='blue', label='Cumulative Score')
    ax.set_title('Cumulative Sentiment Score Chart')
    ax.set_xlabel('Comment Index')
    ax.set_ylabel('Cumulative Score')
    ax.xaxis.set_major_locator(MaxNLocator(integer=True))
    ax.grid(True)
    ax.legend()

    # Save the plot as an image
    chart_path = "cucumber_chart.png"
    plt.savefig(chart_path, format='png', bbox_inches='tight')
    plt.close(fig)
    return chart_path

# Function to create a Word document with a cumulative chart
def generate_word_doc_with_chart(aggressive_comments):
    doc = Document()
    doc.add_heading('Aggressive Comments Report', level=1)

    for comment in aggressive_comments:
        doc.add_heading(f"Post Title: {comment['post_title']}", level=2)
        doc.add_paragraph(f"Comment: {comment['comment_body']}")
        doc.add_paragraph(f"Sentiment Score: {comment['sentiment_score']:.2f}")

    if aggressive_comments:
        chart_path = create_cucumber_chart(aggressive_comments)
        doc.add_page_break()
        doc.add_heading('Cumulative Sentiment Score Chart', level=1)
        doc.add_picture(chart_path, width=docx.shared.Inches(6), height=docx.shared.Inches(4))

    # Save the document to an in-memory buffer
    word_output = io.BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return word_output

# Function to display a pie chart of aggressive vs non-aggressive comments
def plot_aggressive_comments(total_comments, aggressive_count):
    fig, ax = plt.subplots()
    labels = ['Aggressive', 'Non-Aggressive']
    sizes = [aggressive_count, total_comments - aggressive_count]
    colors = ['#FF5733', '#4CAF50']
    explode = (0.1, 0)
    ax.pie(sizes, explode=explode, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90)
    ax.axis('equal')
    st.pyplot(fig)

# Main Streamlit app function
def main():
    st.title("Reddit Behavior Detection")

    search_type = st.radio(
        "Search By:",
        ('Username/Profile Link', 'Subreddit'),
        help="Choose whether to search by Reddit username/profile link or subreddit."
    )

    user_input = st.text_input(
        "Enter Reddit Username, Profile Link, or Subreddit Name:",
        placeholder="e.g., u/username, https://www.reddit.com/user/username, or subreddit_name"
    )

    if st.button("Search"):
        if user_input:
            if search_type == 'Username/Profile Link':
                if "reddit.com" in user_input:
                    username = re.search(r'reddit\.com/user/([a-zA-Z0-9_]+)', user_input).group(1)
                else:
                    username = user_input.strip()

                with st.spinner("Fetching data..."):
                    aggressive_comments, total_comments = fetch_aggressive_comments_by_user(username)

            if total_comments > 0:
                aggressive_count = len(aggressive_comments)
                st.write(f"Total Comments Analyzed: {total_comments}")
                st.write(f"Aggressive Comments Found: {aggressive_count}")

                with stqdm(range(100), desc="Generating Report...", total=100) as progress_bar:
                    for _ in range(100):
                        sleep(0.1)
                        progress_bar.update(1)

                if aggressive_comments:
                    word_file = generate_word_doc_with_chart(aggressive_comments)
                    st.download_button(
                        label="Download Report as Word Document",
                        data=word_file,
                        file_name="aggressive_comments_report.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                else:
                    st.write("No aggressive comments found.")

                plot_aggressive_comments(total_comments, aggressive_count)
            else:
                st.write("No comments available for analysis.")
        else:
            st.error("Please enter a valid input.")

if __name__ == "__main__":
    main()
